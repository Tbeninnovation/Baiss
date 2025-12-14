import os
import pandas as pd
from docx import Document
from baiss_sdk.files.file_reader import FileReader

def extract_chunks(text: str) -> list:
    """
    Placeholder for baiss_sdk.parsers.extract_chunks.
    This simple version splits text by double newlines (paragraphs).
    """
    if not text:
        return []
    return [chunk.strip() for chunk in text.split('\n\n') if chunk.strip()]

class TextDocumentParser:
    """
    A parser for .txt and .docx files that extracts text and tables,
    following the structure of the original PDFParser.
    """

    def __init__(self):
        """Initializes the Text Document Parser."""
        print("Text Document Parser (for .txt and .docx) initialized.")

    def parse(self, file_path: str) -> list:
        """
        Parses a file by dispatching to the correct method based on its extension.

        Args:
            file_path: Path to the file (.txt or .docx).

        Returns:
            A list of dictionaries, each representing a parsed section (the whole document).
        """
        file_path = FileReader.update_file_path(file_path)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at: {file_path}")

        # Get file extension and call the appropriate internal parser
        _, extension = os.path.splitext(file_path.lower())

        if extension == '.txt':
            return self._parse_txt(file_path)
        elif extension == '.docx':
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: '{extension}'. This parser only handles .txt and .docx.")

    def _parse_txt(self, file_path: str) -> list:
        """Handles parsing for .txt files by splitting into 30-line pages."""
        print(f"Parsing TXT file: {os.path.basename(file_path)}")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                full_text = f.read()

            lines = full_text.splitlines()
            pages = [
                "\n".join(lines[i:i+30])
                for i in range(0, len(lines), 30)
                if any(line.strip() for line in lines[i:i+30])
            ]

            parsed_document = []
            for page_num, page_text in enumerate(pages):
                if not page_text.strip():
                    continue

                page_result = {
                    "page_number": page_num + 1,
                    "tags": [],
                    "full_text": page_text.strip(),
                    "chunks": extract_chunks(page_text.strip()),
                    "last_modified": os.path.getmtime(file_path),
                    "tables": [],
                    "images": []
                }
                parsed_document.append(page_result)

            return parsed_document
        except:
            return []

    def _parse_docx(self, file_path: str) -> list:
        """
        Handles parsing for .docx files by splitting into 30-line chunks.
        """
        print(f"Parsing DOCX file with 30-line chunk logic: {os.path.basename(file_path)}")
        doc = Document(file_path)
        
        # Extract all text from paragraphs into a single string
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        lines = full_text.splitlines()
        pages = [
            "\n".join(lines[i:i+30])
            for i in range(0, len(lines), 30)
            if any(line.strip() for line in lines[i:i+30])
        ]

        parsed_document = []
        for page_num, page_text in enumerate(pages):
            if not page_text.strip():
                continue
                
            page_result = {
                "page_number": page_num + 1,
                "tags": [],
                "full_text": page_text.strip(),
                "chunks": extract_chunks(page_text.strip()),
                "last_modified": os.path.getmtime(file_path),
                "tables": [],
                "images": []
            }
            parsed_document.append(page_result)

        # Extract all tables and add them to the last page
        if doc.tables and parsed_document:
            extracted_tables = []
            for table in doc.tables:
                table_data = [[cell.text for cell in row.cells] for row in table.rows]
                extracted_tables.append(table_data)
            
            if extracted_tables:
                parsed_document[-1]["tables"] = extracted_tables
                parsed_document[-1]["tags"].append("table")

        # Handle images tag
        if doc.inline_shapes and parsed_document:
            parsed_document[-1]["tags"].append("image")

        # Fallback if document is empty or only contains tables
        if not parsed_document:
            doc_result = {
                "page_number": 1, 
                "tags": [],
                "full_text": "",
                "chunks": [],
                "tables": [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables],
                "images": []
            }
            if doc_result["tables"]:
                doc_result["tags"].append("table")
            if doc.inline_shapes:
                doc_result["tags"].append("image")
            return [doc_result]

        return parsed_document

    def get_structure(self, parsed_document: list):
        """
        Processes tables from a parsed document and prints them as pandas DataFrames.
        (This is your original method, which now works for both .txt and .docx outputs)
        """
        for page in parsed_document:
            if not page["tables"]:
                continue 

            for i, table_data in enumerate(page["tables"]):
                if not table_data or len(table_data) < 1:
                    print(f"Document contains an empty or invalid table.")
                    continue
                
                try:
                    header = [str(h).strip() if h is not None else '' for h in table_data[0]]
                    data = table_data[1:]
                    df = pd.DataFrame(data, columns=header)
                    
                    print(f"\n--- Table {i+1} from Document ---")
                    print(df)

                except Exception as e:
                    print(f"Error processing table in document: {e}")


if __name__ == "__main__":


    txt_path = ""

    docx_path = ""


    parser = TextDocumentParser()
    import json

    print("\n" + "="*50)
    parsed_txt = parser.parse(txt_path)
    print("\nParsed TXT Output:")
    print(json.dumps(parsed_txt, indent=2))
    parser.get_structure(parsed_txt)
    print("*************************")
    parsed_docx = parser.parse(docx_path)
    print("\nParsed DOCX Output:")
    print(json.dumps(parsed_docx, indent=2))

    parser.get_structure(parsed_docx)